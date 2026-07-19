"""Resume export helpers for PDF and DOCX."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from jinja2 import Environment, FileSystemLoader, select_autoescape
from weasyprint import HTML, CSS

from resume_maker.builder import ResumeData


TEMPLATES_DIR = Path(__file__).parent / "templates"
TEMPLATE_NAME = "resume_template.html"
STYLE_NAME = "resume_style.css"


def _render_resume_html(resume_data: ResumeData) -> str:
    env = Environment(
        loader=FileSystemLoader(str(TEMPLATES_DIR)),
        autoescape=select_autoescape(["html", "xml"]),
    )
    template = env.get_template(TEMPLATE_NAME)
    return template.render(resume=resume_data.to_dict())


def export_to_pdf(resume_data: ResumeData, output_path: str) -> str:
    """Render a resume to PDF using Jinja2 + WeasyPrint."""
    html = _render_resume_html(resume_data)
    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)

    css_path = TEMPLATES_DIR / STYLE_NAME
    stylesheets = [CSS(filename=str(css_path))] if css_path.exists() else None

    HTML(string=html, base_url=str(TEMPLATES_DIR)).write_pdf(
        str(target),
        stylesheets=stylesheets,
    )
    return str(target)


def export_to_docx(resume_data: ResumeData, output_path: str) -> str:
    """Create a clean DOCX resume using python-docx."""
    data = resume_data.to_dict()
    doc = Document()

    doc.add_heading(data.get("name", ""), level=1)

    contact_parts = [
        data.get("email", ""),
        data.get("phone", ""),
        data.get("location", ""),
        data.get("linkedin", ""),
        data.get("github", ""),
    ]
    doc.add_paragraph(" | ".join([part for part in contact_parts if part]))

    doc.add_heading("Objective", level=2)
    doc.add_paragraph(data.get("objective", ""))

    doc.add_heading("Education", level=2)
    for entry in data.get("education", []):
        summary = " - ".join(
            [v for v in [entry.get("degree", ""), entry.get("institution", ""), entry.get("year", "")] if v]
        )
        if summary:
            doc.add_paragraph(summary)
        score = entry.get("score", "")
        if score:
            doc.add_paragraph(f"Score: {score}")

    doc.add_heading("Technical Skills", level=2)
    for category, values in data.get("skills", {}).items():
        if isinstance(values, list):
            value_text = ", ".join(values)
        else:
            value_text = str(values)
        doc.add_paragraph(f"{category}: {value_text}")

    doc.add_heading("Projects", level=2)
    for project in data.get("projects", []):
        header = " - ".join(
            [v for v in [project.get("title", ""), project.get("technologies", "")] if v]
        )
        if header:
            doc.add_paragraph(header)
        if project.get("description"):
            doc.add_paragraph(project["description"])
        for bullet in project.get("bullets", []):
            doc.add_paragraph(str(bullet), style="List Bullet")

    doc.add_heading("Experience", level=2)
    for exp in data.get("experience", []):
        header = " - ".join(
            [v for v in [exp.get("role", ""), exp.get("company", ""), exp.get("duration", "")] if v]
        )
        if header:
            doc.add_paragraph(header)
        for bullet in exp.get("bullets", []):
            doc.add_paragraph(str(bullet), style="List Bullet")

    doc.add_heading("Certifications", level=2)
    for cert in data.get("certifications", []):
        doc.add_paragraph(str(cert), style="List Bullet")

    doc.add_heading("Languages", level=2)
    for language in data.get("languages", []):
        doc.add_paragraph(str(language), style="List Bullet")

    target = Path(output_path)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    return str(target)
